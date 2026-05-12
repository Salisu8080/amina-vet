"""
build_docx.py — ABU SPGS-compliant dissertation .docx builder
Student : Balarabe, Amina Abdullahi | P24VTVH8002
Degree  : MSc Veterinary Public Health
Output  : docs/dissertation_draft.docx

Usage:
    python build_docx.py

Chapter .md files expected in docs/:
    Preliminary_Pages.md      (reference only — not rendered directly)
    Chapter1_Introduction.md
    Chapter2_Literature_Review.md
    Chapter3_Materials_and_Methods.md
    Chapter4_Results.md
    Chapter5_Discussion.md
    Chapter6_Conclusion.md
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
DOCS = BASE / "docs"
OUT  = DOCS / "dissertation_draft.docx"

# ── ABU SPGS constants ────────────────────────────────────────────────────────
FONT_NAME   = "Times New Roman"
FONT_SIZE   = Pt(12)
MARGIN_TOP  = Inches(1)
MARGIN_BOT  = Inches(1)
MARGIN_LEFT = Inches(1.5)   # binding margin
MARGIN_RGT  = Inches(1)

# ── Student / dissertation constants ─────────────────────────────────────────
TITLE_UPPER = (
    "PREVALENCE OF MULTIDRUG-RESISTANT ESCHERICHIA COLI BEARING "
    "CLASS I AND CLASS II INTEGRONS IN POULTRY MANURE FROM "
    "VENDORS IN ZARIA METROPOLIS"
)
STUDENT_NAME    = "Balarabe, Amina Abdullahi"
STUDENT_DEGREE  = "B.Sc. Microbiology (ABU Zaria, [Year of First Degree])"
REG_NUMBER      = "P24VTVH8002"
DEPT_LINE1      = "DEPARTMENT OF VETERINARY PUBLIC HEALTH AND PREVENTIVE MEDICINE,"
DEPT_LINE2      = "FACULTY OF VETERINARY MEDICINE,"
DEPT_LINE3      = "AHMADU BELLO UNIVERSITY, ZARIA, NIGERIA."
DEPT_SHORT      = "Department of Veterinary Public Health and Preventive Medicine"
DEGREE_AWARD    = "MASTER OF SCIENCE (MSc) IN VETERINARY PUBLIC HEALTH"
SUPERVISOR_CH   = "Prof. Junaidu Kabir"
SUPERVISOR_MEM  = "Dr. U.U. Bashir"
HOD_NAME        = "[Head of Department — Full Name and Rank]"
PG_COORD        = "[Postgraduate Coordinator — Full Name and Rank]"
PROVOST_NAME    = "Prof. O. O. Okubanjo"


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def set_margins(section):
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOT
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RGT


def set_font(run, bold=False, italic=False, size=None):
    run.font.name   = FONT_NAME
    run.font.size   = size or FONT_SIZE
    run.font.bold   = bold
    run.font.italic = italic


def apply_paragraph_format(para, spacing="double",
                            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=Pt(0), space_after=Pt(6)):
    pf = para.paragraph_format
    pf.alignment    = align
    pf.space_before = space_before
    pf.space_after  = space_after
    if spacing == "double":
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif spacing == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif spacing == "1_5":
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_para(doc, text="", bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing="double",
             space_before=Pt(0), space_after=Pt(6), size=None):
    para = doc.add_paragraph()
    apply_paragraph_format(para, spacing=spacing, align=align,
                           space_before=space_before, space_after=space_after)
    if text:
        run = para.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size)
    return para


def add_heading1(doc, text, spacing="double"):
    """Chapter title: UPPER CASE BOLD, 14pt, centred — picked up by TOC."""
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 1"]
    apply_paragraph_format(para, spacing=spacing,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(12), space_after=Pt(6))
    run = para.add_run(text.upper())
    set_font(run, bold=True, size=Pt(14))
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_heading2(doc, text, spacing="double"):
    """Section heading: Title Case Bold, left-aligned."""
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 2"]
    apply_paragraph_format(para, spacing=spacing,
                           align=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=Pt(12), space_after=Pt(6))
    run = para.add_run(text)
    set_font(run, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_heading3(doc, text, spacing="double"):
    """Sub-section: Sentence case bold, left-aligned."""
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 3"]
    apply_paragraph_format(para, spacing=spacing,
                           align=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=Pt(6), space_after=Pt(3))
    run = para.add_run(text)
    set_font(run, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_heading4(doc, text, spacing="double"):
    """Sub-sub-section: bold italic, left-aligned."""
    para = doc.add_paragraph()
    apply_paragraph_format(para, spacing=spacing,
                           align=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=Pt(6), space_after=Pt(3))
    run = para.add_run(text)
    set_font(run, bold=True, italic=True)
    return para


def add_section_break(doc, break_type=WD_SECTION.NEW_PAGE):
    new_section = doc.add_section(break_type)
    set_margins(new_section)
    return new_section


def set_page_number_format(section, fmt="lowerRoman", start=1):
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def add_page_number_footer(doc, section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    run = fp.add_run()
    set_font(run)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def insert_toc(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run = para.add_run()
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)
    placeholder = doc.add_paragraph(
        "[Right-click here → 'Update Field' to generate Table of Contents]"
    )
    apply_paragraph_format(placeholder, spacing="single",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(placeholder.runs[0], italic=True)


# ── Inline markdown parser ─────────────────────────────────────────────────────

def add_inline_text(para, text, base_bold=False, base_italic=False, size=None):
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            set_font(run, bold=base_bold, italic=base_italic, size=size)
        if m.group(2):
            run = para.add_run(m.group(2))
            set_font(run, bold=True, italic=True, size=size)
        elif m.group(3):
            run = para.add_run(m.group(3))
            set_font(run, bold=True, italic=base_italic, size=size)
        elif m.group(4):
            run = para.add_run(m.group(4))
            set_font(run, bold=base_bold, italic=True, size=size)
        elif m.group(5):
            run = para.add_run(m.group(5))
            set_font(run, bold=base_bold, italic=True, size=size)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        set_font(run, bold=base_bold, italic=base_italic, size=size)


def add_rich_para(doc, text, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing="double",
                  space_before=Pt(0), space_after=Pt(6), size=None):
    para = doc.add_paragraph()
    apply_paragraph_format(para, spacing=spacing, align=align,
                           space_before=space_before, space_after=space_after)
    add_inline_text(para, text, base_bold=bold, base_italic=italic, size=size)
    return para


def add_md_table(doc, lines, spacing="double"):
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r"^\|[-:| ]+\|$", line):
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line[1:-1].split("|")]
            rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= ncols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            apply_paragraph_format(para, spacing="single",
                                   space_before=Pt(2), space_after=Pt(2))
            add_inline_text(para, cell_text, base_bold=(ri == 0))
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Markdown → docx renderer ──────────────────────────────────────────────────

def render_markdown(doc, md_text, default_spacing="double"):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() in ("&nbsp;", ""):
            i += 1
            continue

        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        if line.startswith(">"):
            text = re.sub(r"^>+\s*", "", line)
            while i + 1 < len(lines) and lines[i + 1].startswith(">"):
                i += 1
                text += " " + re.sub(r"^>+\s*", "", lines[i])
            if text.strip():
                p = doc.add_paragraph()
                apply_paragraph_format(p, spacing="single",
                                       align=WD_ALIGN_PARAGRAPH.LEFT,
                                       space_before=Pt(3), space_after=Pt(3))
                p.paragraph_format.left_indent = Inches(0.5)
                add_inline_text(p, f"Note: {text.strip()}", base_italic=True)
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(11)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_md_table(doc, table_lines, spacing=default_spacing)
            continue

        m = re.match(r"^#{1}\s+(.+)$", line)
        if m and not re.match(r"^#{2,}", line):
            add_heading1(doc, m.group(1).strip(), spacing=default_spacing)
            i += 1
            continue

        m = re.match(r"^#{2}\s+(.+)$", line)
        if m and not re.match(r"^#{3,}", line):
            add_heading2(doc, m.group(1).strip(), spacing=default_spacing)
            i += 1
            continue

        m = re.match(r"^#{3}\s+(.+)$", line)
        if m and not re.match(r"^#{4,}", line):
            add_heading3(doc, m.group(1).strip(), spacing=default_spacing)
            i += 1
            continue

        m = re.match(r"^#{4}\s+(.+)$", line)
        if m:
            add_heading4(doc, m.group(1).strip(), spacing=default_spacing)
            i += 1
            continue

        m = re.match(r"^[-*+]\s+(.+)$", line.strip())
        if m:
            p = doc.add_paragraph(style="List Bullet")
            apply_paragraph_format(p, spacing="single",
                                   space_before=Pt(2), space_after=Pt(2))
            add_inline_text(p, m.group(1))
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
            i += 1
            continue

        m = re.match(r"^\d+\.\s+(.+)$", line.strip())
        if m:
            p = doc.add_paragraph(style="List Number")
            apply_paragraph_format(p, spacing="single",
                                   space_before=Pt(2), space_after=Pt(2))
            add_inline_text(p, m.group(1))
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
            i += 1
            continue

        # Reference line (hanging indent)
        if re.match(r"^- .{5,}", line) and (
            "doi" in line.lower() or re.search(r"\(\d{4}\)", line)
        ):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent       = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)
            pf.space_before      = Pt(0)
            pf.space_after       = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_inline_text(p, line[2:].strip())
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
            i += 1
            continue

        # Normal paragraph
        if line.strip():
            para_lines = [line.strip()]
            while (
                i + 1 < len(lines)
                and lines[i + 1].strip()
                and not lines[i + 1].startswith("#")
                and not lines[i + 1].startswith("|")
                and not lines[i + 1].startswith(">")
                and not lines[i + 1].startswith("```")
                and not re.match(r"^---+$", lines[i + 1].strip())
                and not re.match(r"^[-*+]\s", lines[i + 1].strip())
                and not re.match(r"^\d+\.\s", lines[i + 1].strip())
            ):
                i += 1
                para_lines.append(lines[i].strip())
            full_text = " ".join(para_lines)
            if full_text.strip():
                add_rich_para(doc, full_text, spacing=default_spacing)

        i += 1


# ══════════════════════════════════════════════════════════════════════════════
# PRELIMINARY PAGE BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def build_cover_page(doc):
    for _ in range(6):
        add_para(doc, spacing="single", space_after=Pt(0))

    add_para(doc, "[AHMADU BELLO UNIVERSITY LOGO — INSERT HERE]",
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single",
             italic=True, space_after=Pt(24))

    add_para(doc, TITLE_UPPER, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             spacing="single", space_after=Pt(18))
    add_para(doc, "BY", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(18))
    add_para(doc, STUDENT_NAME, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(24))

    for line in [DEPT_LINE1.rstrip(","), DEPT_LINE2.rstrip(","), "AHMADU BELLO UNIVERSITY, ZARIA"]:
        add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 spacing="single", space_after=Pt(2))

    add_para(doc, "", spacing="single", space_after=Pt(24))
    add_para(doc, "[MONTH, YEAR OF CORRECTION CERTIFICATION]", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(0))


def build_title_page(doc):
    for _ in range(4):
        add_para(doc, spacing="single", space_after=Pt(0))

    add_para(doc, TITLE_UPPER, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(18))
    add_para(doc, "BY", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(12))
    add_para(doc, STUDENT_NAME, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(2))
    add_para(doc, f"({STUDENT_DEGREE})", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(2))
    add_para(doc, REG_NUMBER, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(18))

    submission = (
        f"A DISSERTATION SUBMITTED TO THE POSTGRADUATE COLLEGE, "
        f"AHMADU BELLO UNIVERSITY IN PARTIAL FULFILMENT OF THE REQUIREMENTS "
        f"FOR THE AWARD OF {DEGREE_AWARD}"
    )
    add_para(doc, submission, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(18))

    for line in [DEPT_LINE1, DEPT_LINE2, DEPT_LINE3]:
        add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 spacing="single", space_after=Pt(2))

    add_para(doc, "", spacing="single", space_after=Pt(18))
    add_para(doc, "[MONTH], [YEAR]", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single", space_after=Pt(0))


def build_declaration(doc):
    add_heading2(doc, "DECLARATION", spacing="single")
    text = (
        f'I declare that the work in this Dissertation entitled "{TITLE_UPPER}" '
        f"was carried out by me in the {DEPT_SHORT}. The information derived from "
        "the literature was duly acknowledged in the text and a list of references "
        "provided. No part of this Dissertation was previously presented for another "
        "degree or diploma at this or any other Institution."
    )
    add_rich_para(doc, text, spacing="single")
    add_para(doc, spacing="single", space_after=Pt(24))

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    row0 = [STUDENT_NAME, "___________________", "___________________"]
    row1 = ["Name of Student", "Signature", "Date"]
    for ci, (h, l) in enumerate(zip(row0, row1)):
        for ri, val in enumerate([h, l]):
            cell = table.cell(ri, ci)
            cell.text = val
            p = cell.paragraphs[0]
            apply_paragraph_format(p, spacing="single",
                                   space_before=Pt(2), space_after=Pt(2))
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE


def build_certification(doc):
    add_heading2(doc, "CERTIFICATION", spacing="single")
    text = (
        f'This Dissertation entitled "{TITLE_UPPER}" by {STUDENT_NAME} meets the '
        f"regulations governing the award of the degree of {DEGREE_AWARD} of the "
        "Ahmadu Bello University, and is approved for its contribution to knowledge "
        "and literary presentation."
    )
    add_rich_para(doc, text, spacing="single")
    add_para(doc, spacing="single", space_after=Pt(12))

    rows_data = [
        ["Role", "Name", "Signature", "Date"],
        ["Chairman, Supervisory Committee", SUPERVISOR_CH,  "____________", "______"],
        ["Member, Supervisory Committee",  SUPERVISOR_MEM, "____________", "______"],
        ["Head of Department",             HOD_NAME,       "____________", "______"],
        ["Postgraduate Coordinator",       PG_COORD,       "____________", "______"],
        ["Provost, Postgraduate College",  PROVOST_NAME,   "____________", "______"],
    ]
    table = doc.add_table(rows=len(rows_data), cols=4)
    table.style = "Table Grid"
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = val
            p = cell.paragraphs[0]
            apply_paragraph_format(p, spacing="single",
                                   space_before=Pt(2), space_after=Pt(2))
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
                if ri == 0:
                    run.font.bold = True


def build_dedication(doc):
    add_heading2(doc, "DEDICATION", spacing="single")
    add_para(
        doc,
        "This work is dedicated to Almighty Allah, the Most Gracious and Most Merciful, "
        "and to my beloved family whose prayers, love, and unwavering support sustained "
        "me throughout this academic journey.",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single",
    )


def build_acknowledgement(doc):
    add_heading2(doc, "ACKNOWLEDGEMENT", spacing="single")
    paras = [
        "All praises and thanks are due to Almighty Allah (Subhanahu wa Ta'ala), the "
        "Lord of the worlds, for granting me the knowledge, patience, and strength to "
        "complete this work.",

        f"I wish to express my profound gratitude to my supervisory committee: "
        f"{SUPERVISOR_CH} (Chairman) and {SUPERVISOR_MEM} (Member), for their "
        "exceptional guidance, invaluable advice, and constructive criticisms throughout "
        "the period of this study. Their commitment to excellence greatly shaped the "
        "quality of this work.",

        "My sincere appreciation goes to the Head of Department and all academic and "
        f"support staff of the {DEPT_SHORT}, Faculty of Veterinary Medicine, "
        "Ahmadu Bello University, Zaria, for their assistance and the use of departmental "
        "facilities.",

        "Special thanks to the laboratory staff of the Bacterial Zoonoses Laboratory "
        "for their technical support during sample processing, bacterial isolation, "
        "biochemical characterisation, antimicrobial susceptibility testing, and "
        "molecular analysis.",

        "Finally, my heartfelt thanks go to my family for their endless support, "
        "encouragement, and prayers throughout my academic journey. May Almighty Allah "
        "reward them abundantly.",
    ]
    for p_text in paras:
        add_rich_para(doc, p_text, spacing="single")


def build_abstract(doc):
    add_heading2(doc, "ABSTRACT", spacing="single")
    abstract_path = DOCS / "Abstract.md"
    if abstract_path.exists():
        raw = abstract_path.read_text(encoding="utf-8")
        # Extract only the body: lines that are not headings, metadata, dividers,
        # keyword lines, or word-count notes.
        body_lines = []
        for line in raw.splitlines():
            s = line.strip()
            if not s:
                continue
            if s.startswith("#"):
                continue
            if s.startswith("**Title:") or s.startswith("**Student:"):
                continue
            if re.match(r"^---+$", s):
                continue
            if s.startswith("**Keywords:") or s.startswith("*(Word count"):
                continue
            body_lines.append(s)
        body = " ".join(body_lines).strip()
        if body:
            add_rich_para(doc, body, spacing="single")
        else:
            add_rich_para(
                doc,
                "[ABSTRACT — TO BE WRITTEN. Maximum 500 words. "
                "Write content in docs/Abstract.md then rebuild.]",
                italic=True, spacing="single",
            )
    else:
        add_rich_para(
            doc,
            "[ABSTRACT — TO BE WRITTEN AFTER ALL CHAPTERS ARE COMPLETE. "
            "Maximum 500 words. Single spacing. Times New Roman 12pt, justified. "
            "Write content in docs/Abstract.md then rebuild.]",
            italic=True, spacing="single",
        )
    add_para(doc, spacing="single", space_after=Pt(6))
    kw = doc.add_paragraph()
    apply_paragraph_format(kw, spacing="single",
                           space_before=Pt(6), space_after=Pt(0))
    r1 = kw.add_run("Keywords: ")
    set_font(r1, bold=True)
    r2 = kw.add_run(
        "Escherichia coli, multidrug resistance, integrons, poultry manure, "
        "antimicrobial resistance, Zaria, Nigeria, Class I integron, "
        "Class II integron, MAR index, One Health, veterinary public health"
    )
    set_font(r2)


def build_toc_page(doc):
    add_heading2(doc, "TABLE OF CONTENTS", spacing="single")
    insert_toc(doc)


def build_list_of_tables(doc):
    add_heading2(doc, "LIST OF TABLES", spacing="single")
    tables_data = [
        ("Table", "Title", "Page"),
        ("3.1",   "Antibiotic discs used for antimicrobial susceptibility testing (CLSI, 2024)", "—"),
        ("3.2",   "Primers used for PCR amplification (16S rRNA, Class I and Class II integrons)", "—"),
        ("4.1",   "Distribution of *E. coli* isolates by sampling location in Zaria Metropolis", "—"),
        ("4.2",   "Antimicrobial susceptibility patterns of *E. coli* isolates (n=37)", "—"),
        ("4.3",   "Multiple Antibiotic Resistance (MAR) index of *E. coli* isolates", "—"),
        ("4.4",   "Distinct antimicrobial resistance patterns among MDR *E. coli* isolates", "—"),
        ("4.5",   "Prevalence of Class I and Class II integrons among *E. coli* isolates by location", "—"),
    ]
    add_md_table(doc, [f"| {r[0]} | {r[1]} | {r[2]} |" for r in tables_data])


def build_list_of_plates(doc):
    add_heading2(doc, "LIST OF PLATES", spacing="single")
    plates_data = [
        ("Plate", "Description", "Page"),
        ("4.1",   "Eosin Methylene Blue (EMB) agar showing metallic green sheen colonies of *E. coli*", "—"),
        ("4.2",   "Agarose gel electrophoresis showing PCR amplification of 16S rRNA gene (expected band: ~1,500 bp) confirming *E. coli* isolates", "—"),
        ("4.3",   "Agarose gel electrophoresis showing PCR amplification of Class I integron (1,100 bp) and Class II integron (233 bp) gene fragments in *E. coli* isolates", "—"),
    ]
    add_md_table(doc, [f"| {r[0]} | {r[1]} | {r[2]} |" for r in plates_data])


def build_list_of_abbreviations(doc):
    add_heading2(doc, "LIST OF ABBREVIATIONS", spacing="single")
    abbrevs = [
        ("Abbreviation", "Full Form"),
        ("ABU",   "Ahmadu Bello University"),
        ("AMR",   "Antimicrobial Resistance"),
        ("ARB",   "Antibiotic Resistant Bacteria"),
        ("ARG",   "Antibiotic Resistance Gene"),
        ("AST",   "Antibiotic Susceptibility Testing"),
        ("AZM",   "Azithromycin"),
        ("BPW",   "Buffered Peptone Water"),
        ("C",     "Chloramphenicol"),
        ("CIP",   "Ciprofloxacin"),
        ("CLSI",  "Clinical Laboratory Standards Institute"),
        ("CN",    "Gentamicin"),
        ("CRO",   "Ceftriaxone"),
        ("DNA",   "Deoxyribonucleic Acid"),
        ("EMB",   "Eosin Methylene Blue"),
        ("ESBL",  "Extended Spectrum Beta-Lactamase"),
        ("FAO",   "Food and Agriculture Organization of the United Nations"),
        ("HGT",   "Horizontal Gene Transfer"),
        ("LGA",   "Local Government Area"),
        ("MAR",   "Multiple Antibiotic Resistance"),
        ("MDR",   "Multidrug Resistant/Resistance"),
        ("MEM",   "Meropenem"),
        ("MGE",   "Mobile Genetic Element"),
        ("MR-VP", "Methyl Red–Voges Proskauer"),
        ("NA",    "Nalidixic Acid"),
        ("PCR",   "Polymerase Chain Reaction"),
        ("rRNA",  "Ribosomal Ribonucleic Acid"),
        ("SIM",   "Sulfate Indole Motility"),
        ("SPSS",  "Statistical Package for Social Sciences"),
        ("SXT",   "Cotrimoxazole (Trimethoprim-Sulfamethoxazole)"),
        ("TAE",   "Tris Acetate EDTA"),
        ("TE",    "Tetracycline"),
        ("TSB",   "Tryptone Soy Broth"),
        ("TSI",   "Triple Sugar Iron"),
        ("UV",    "Ultraviolet"),
        ("VPH",   "Veterinary Public Health"),
        ("WHO",   "World Health Organization"),
        ("WOAH",  "World Organisation for Animal Health"),
    ]
    add_md_table(doc, [f"| {r[0]} | {r[1]} |" for r in abbrevs])


# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES AND APPENDICES BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_references_and_appendices(doc):
    ref_app_path = DOCS / "References_Appendices.md"
    if not ref_app_path.exists():
        add_heading1(doc, "REFERENCES", spacing="double")
        add_rich_para(doc, "[References not yet compiled.]", italic=True)
        return

    text = ref_app_path.read_text(encoding="utf-8")

    # Split at ## APPENDICES boundary
    parts = re.split(r"^## APPENDICES", text, flags=re.MULTILINE)
    ref_section  = parts[0]
    app_section  = parts[1] if len(parts) > 1 else ""

    # Strip the ## REFERENCES heading and any preamble notes
    ref_body_parts = re.split(r"^## REFERENCES", ref_section, flags=re.MULTILINE)
    ref_body = ref_body_parts[1] if len(ref_body_parts) > 1 else ref_section

    # Parse individual reference entries (blank-line separated paragraphs)
    entries = []
    current = []
    for line in ref_body.splitlines():
        s = line.strip()
        if not s or s.startswith(">") or re.match(r"^---+$", s) or s.startswith("*("):
            if current:
                entries.append(" ".join(current))
                current = []
        else:
            current.append(s)
    if current:
        entries.append(" ".join(current))

    # Render REFERENCES section
    add_heading1(doc, "REFERENCES", spacing="double")
    for entry in entries:
        if entry.strip():
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.left_indent       = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)
            pf.space_before      = Pt(0)
            pf.space_after       = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_inline_text(p, entry)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE

    # Render APPENDICES section
    if app_section.strip():
        doc.add_page_break()
        add_heading1(doc, "APPENDICES", spacing="double")
        clean = re.sub(r"^[\s\-]+", "", app_section, count=3)
        render_markdown(doc, clean, default_spacing="single")


# ══════════════════════════════════════════════════════════════════════════════
# STYLES INITIALISER
# ══════════════════════════════════════════════════════════════════════════════

def init_styles(doc):
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = FONT_SIZE


# ══════════════════════════════════════════════════════════════════════════════
# MAIN BUILD
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    init_styles(doc)

    # ── Cover page (no numbering) ────────────────────────────────────────────
    sec0 = doc.sections[0]
    set_margins(sec0)
    build_cover_page(doc)

    # ── Fly leaf ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_para(doc, "[Intentionally blank — fly leaf]",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing="single")

    # ── Preliminary pages: Roman numerals i, ii, iii … ───────────────────────
    sec_prelim = add_section_break(doc, WD_SECTION.NEW_PAGE)
    set_page_number_format(sec_prelim, fmt="lowerRoman", start=1)
    add_page_number_footer(doc, sec_prelim)

    build_title_page(doc);       doc.add_page_break()
    build_declaration(doc);      doc.add_page_break()
    build_certification(doc);    doc.add_page_break()
    build_dedication(doc);       doc.add_page_break()
    build_acknowledgement(doc);  doc.add_page_break()
    build_abstract(doc);         doc.add_page_break()
    build_toc_page(doc);         doc.add_page_break()
    build_list_of_tables(doc);   doc.add_page_break()
    build_list_of_plates(doc);   doc.add_page_break()
    build_list_of_abbreviations(doc)

    # ── Chapter pages: Arabic numerals 1, 2, 3 … ─────────────────────────────
    sec_chapters = add_section_break(doc, WD_SECTION.NEW_PAGE)
    set_page_number_format(sec_chapters, fmt="decimal", start=1)
    add_page_number_footer(doc, sec_chapters)

    chapter_files = [
        DOCS / "Chapter1_Introduction.md",
        DOCS / "Chapter2_Literature_Review.md",
        DOCS / "Chapter3_Methodology.md",
        DOCS / "Chapter4_Results.md",
        DOCS / "Chapter5_Discussion.md",
        DOCS / "Chapter6_Conclusion.md",
    ]

    for fp in chapter_files:
        if fp.exists():
            text = fp.read_text(encoding="utf-8")
            render_markdown(doc, text, default_spacing="double")
            doc.add_page_break()
        else:
            add_heading1(doc, f"[{fp.stem} — NOT YET WRITTEN]")
            doc.add_page_break()

    # ── References and Appendices ─────────────────────────────────────────────
    build_references_and_appendices(doc)

    # ── Save ─────────────────────────────────────────────────────────────────
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved : {OUT}")
    print(f"Size  : {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
